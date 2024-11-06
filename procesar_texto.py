# -*- coding: utf-8 -*-
"""
Created on Wed Nov  6 11:23:52 2024

@author: jveraz
"""

from docx import Document
import re

# Cargar el documento original
input_path = "Danza Chunchada de Paucartambo_ historia y fundación [Documental].docx"
doc = Document(input_path)

# Crear un nuevo documento para guardar el texto limpio
output_doc = Document()

# Expresión regular para detectar timestamps como "00:00"
timestamp_pattern = re.compile(r"^\d{2}:\d{2},")

# Procesar el texto eliminando timestamps y añadiéndolo al nuevo documento
for para in doc.paragraphs:
    text = para.text.strip()
    if timestamp_pattern.match(text):
        # Eliminar el timestamp y cualquier espacio extra
        processed_text = text[6:].strip()
    else:
        processed_text = text

    # Añadir el texto limpio al nuevo documento, evitando líneas vacías
    if processed_text:
        output_doc.add_paragraph(processed_text)

# Guardar el documento procesado
output_doc.save("Danza_Chunchada_Transcript_Cleaned.docx")
print("Documento limpio guardado como 'Danza_Chunchada_Transcript_Cleaned.docx'")
